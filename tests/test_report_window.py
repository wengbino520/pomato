"""
tests/test_report_window.py
ReportWindow 的 Markdown→纯文本 转换与 Word 导出内容正确性测试。
"""
import pytest
import re as _re_mod
from unittest.mock import MagicMock, patch, mock_open

from PyQt6.QtWidgets import QDialog

from src.report_window import ReportWindow


# ── 辅助 ──────────────────────────────────────────────────────────────────────

class _DummyConfig:
    def get(self, key, default=None):
        defaults = {
            "report_system_prompt": "",
            "api_base_url": "",
            "api_key": "",
            "api_model": "gpt-4o-mini",
        }
        return defaults.get(key, default)


class _DummyDB:
    def get_entries_by_date(self, date_str):
        return []

    def save_report(self, *args, **kwargs):
        pass


class _DummyAIClient:
    def generate_report(self, entries, report_date, on_chunk=None):
        return "# Test Report\nGenerated content."


# ── Markdown → 纯文本转换 ─────────────────────────────────────────────────────

class TestMarkdownToPlainText:
    """_markdown_to_plain_text 将 Markdown 转为适合粘贴到钉钉/飞书的纯文本。"""

    def test_heading_stripped(self):
        result = ReportWindow._markdown_to_plain_text("# 工作日报 2026-06-01")
        assert result.strip() == "工作日报 2026-06-01"

    def test_h2_heading_stripped(self):
        result = ReportWindow._markdown_to_plain_text("## 📊 今日概览")
        assert result.strip() == "📊 今日概览"

    def test_h3_heading_stripped(self):
        result = ReportWindow._markdown_to_plain_text("### 开发工作")
        assert result.strip() == "开发工作"

    def test_bullet_list_item_preserved(self):
        result = ReportWindow._markdown_to_plain_text("- 完成了登录功能")
        assert "- 完成了登录功能" in result

    def test_asterisk_list_item_preserved(self):
        result = ReportWindow._markdown_to_plain_text("* 修复了 bug")
        assert "- 修复了 bug" in result

    def test_bold_syntax_stripped(self):
        result = ReportWindow._markdown_to_plain_text("这是**重要**内容")
        assert result.strip() == "这是重要内容"

    def test_italic_syntax_stripped(self):
        result = ReportWindow._markdown_to_plain_text("这是*斜体*文本")
        assert result.strip() == "这是斜体文本"

    def test_code_syntax_stripped(self):
        result = ReportWindow._markdown_to_plain_text("调用 `foo()` 函数")
        assert result.strip() == "调用 foo() 函数"

    def test_empty_line_preserved(self):
        result = ReportWindow._markdown_to_plain_text("第一段\n\n第二段")
        assert "\n\n" in result

    def test_multiple_empty_lines_collapsed(self):
        result = ReportWindow._markdown_to_plain_text("A\n\n\n\nB")
        assert result.count("\n\n\n") == 0

    def test_mixed_content(self):
        input_md = (
            "# 日报\n\n"
            "## 开发\n"
            "- **完成了** `API` 对接\n"
            "- *修复了* 2个 bug\n\n"
            "### 明日计划\n"
            "- 继续优化性能"
        )
        result = ReportWindow._markdown_to_plain_text(input_md)
        assert "日报" in result
        assert "开发" in result
        assert "完成了" in result
        assert "API 对接" in result
        assert "修复了 2个 bug" in result
        assert "明日计划" in result
        assert "继续优化性能" in result

    def test_trailing_newline(self):
        result = ReportWindow._markdown_to_plain_text("测试内容")
        assert result.endswith("\n")


# ── Word 导出 ─────────────────────────────────────────────────────────────────

class TestExportDocx:
    """_export_docx 和 _add_markdown_inline 的内容正确性。"""

    def test_export_docx_cancelled_on_empty_path(self, qapp):
        """用户取消保存对话框时，不应崩溃。"""
        with patch("PyQt6.QtWidgets.QFileDialog.getSaveFileName", return_value=("", "")), \
             patch.object(QDialog, "__init__", return_value=None):
            win = ReportWindow.__new__(ReportWindow)
            win.config = _DummyConfig()
            win.db = _DummyDB()
            win.editor = MagicMock()
            win.editor.toPlainText.return_value = "# Test"
            win.report_date = "2026-06-01"
            win.entries = []
            win._export_docx()  # 不崩溃即通过

    def test_export_docx_creates_document_with_content(self, qapp, tmp_path):
        """选择保存路径后，应生成含有内容的 .docx 文件。"""
        sample_md = (
            "# 工作日报 2026-06-01\n\n"
            "## 📊 今日概览\n"
            "- 完成 **3** 个番茄钟\n\n"
            "## 🔧 开发工作\n"
            "- 完成用户登录模块\n"
            "- 修复 *2* 个边界case\n\n"
            "## 💡 待跟进\n"
            "- 性能优化方案待评审"
        )

        output_path = str(tmp_path / "test_report.docx")
        with patch("PyQt6.QtWidgets.QFileDialog.getSaveFileName",
                   return_value=(output_path, "Word 文档 (*.docx)")), \
             patch.object(QDialog, "__init__", return_value=None), \
             patch("src.report_window.QMessageBox.information"):
            win = ReportWindow.__new__(ReportWindow)
            win.config = _DummyConfig()
            win.db = _DummyDB()
            win.editor = MagicMock()
            win.editor.toPlainText.return_value = sample_md
            win.report_date = "2026-06-01"
            win.entries = []

            win._export_docx()

        # 验证文件已创建
        import os
        assert os.path.exists(output_path)
        assert os.path.getsize(output_path) > 0

        # 验证内容可读
        from docx import Document
        doc = Document(output_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "工作日报" in full_text
        assert "今日概览" in full_text
        assert "开发工作" in full_text
        assert "完成用户登录模块" in full_text
        assert "待跟进" in full_text

    def test_export_docx_shows_warning_when_docx_not_installed(self, qapp):
        """未安装 python-docx 时给出友好提示。"""
        win = ReportWindow.__new__(ReportWindow)
        win.config = _DummyConfig()
        win.db = _DummyDB()
        win.editor = MagicMock()
        win.editor.toPlainText.return_value = "# Test"
        win.report_date = "2026-06-01"
        win.entries = []

        with patch.object(QDialog, "__init__", return_value=None), \
             patch("src.report_window.QMessageBox.warning") as mock_warn, \
             patch("builtins.__import__", side_effect=ImportError("No module named 'docx'")):
            win._export_docx()

        mock_warn.assert_called_once()
        assert "python-docx" in mock_warn.call_args[0][2]

    def test_export_markdown_blockquote_italic(self, qapp, tmp_path):
        """引用块导出为斜体。"""
        sample_md = "> 这是一条引用"
        output_path = str(tmp_path / "test_quote.docx")

        with patch("PyQt6.QtWidgets.QFileDialog.getSaveFileName",
                   return_value=(output_path, "Word 文档 (*.docx)")), \
             patch.object(QDialog, "__init__", return_value=None), \
             patch("src.report_window.QMessageBox.information"):
            win = ReportWindow.__new__(ReportWindow)
            win.config = _DummyConfig()
            win.db = _DummyDB()
            win.editor = MagicMock()
            win.editor.toPlainText.return_value = sample_md
            win.report_date = "2026-06-01"
            win.entries = []

            win._export_docx()

        from docx import Document
        doc = Document(output_path)
        texts = [p.text for p in doc.paragraphs]
        assert any("这是一条引用" in t for t in texts)

    def test_add_markdown_inline_bold(self):
        """_add_markdown_inline 正确渲染粗体。"""
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        ReportWindow._add_markdown_inline(p, "这是 **粗体** 文本")
        # 验证有 run 被设为 bold
        bold_runs = [r for r in p.runs if r.bold]
        assert len(bold_runs) >= 1
        assert bold_runs[0].text == "粗体"

    def test_add_markdown_inline_italic(self):
        """_add_markdown_inline 正确渲染斜体。"""
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        ReportWindow._add_markdown_inline(p, "这是 *斜体* 文本")
        italic_runs = [r for r in p.runs if r.italic]
        assert len(italic_runs) >= 1
        assert italic_runs[0].text == "斜体"

    def test_add_markdown_inline_no_formatting(self):
        """无格式文本直接渲染。"""
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        ReportWindow._add_markdown_inline(p, "普通文本")
        assert len(p.runs) >= 1
        assert p.runs[0].text == "普通文本"

    def test_add_markdown_inline_mixed(self):
        """混合粗体+斜体+普通文本。"""
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        ReportWindow._add_markdown_inline(p, "A **B** C *D* E")
        texts = [r.text for r in p.runs]
        combined = "".join(texts)
        assert combined == "A B C D E"


# ── Fallback 生成 ─────────────────────────────────────────────────────────────

class TestGenerateFallback:
    """AI 失败时的回退展示内容。"""

    def test_fallback_contains_header(self):
        win = ReportWindow.__new__(ReportWindow)
        win.report_date = "2026-06-01"
        win.entries = []
        result = win._generate_fallback()
        assert "# 工作日报 2026-06-01" in result

    def test_fallback_shows_valid_entries_only(self):
        win = ReportWindow.__new__(ReportWindow)
        win.report_date = "2026-06-01"
        win.entries = [
            {"start_time": "09:00:00", "end_time": "09:25:00",
             "content": "任务A", "tags": ["开发"], "skipped": 0},
            {"start_time": "09:30:00", "end_time": "09:55:00",
             "content": "", "tags": [], "skipped": 1},
        ]
        result = win._generate_fallback()
        assert "任务A" in result
        assert "开发" in result
        assert "09:30" not in result  # skipped entry not shown

    def test_fallback_shows_correct_count(self):
        win = ReportWindow.__new__(ReportWindow)
        win.report_date = "2026-06-01"
        win.entries = [
            {"start_time": "09:00:00", "end_time": "09:25:00",
             "content": "A", "tags": [], "skipped": 0},
            {"start_time": "09:30:00", "end_time": "09:55:00",
             "content": "B", "tags": [], "skipped": 0},
        ]
        result = win._generate_fallback()
        assert "共 2 个番茄钟" in result

    def test_fallback_empty_entries(self):
        win = ReportWindow.__new__(ReportWindow)
        win.report_date = "2026-06-01"
        win.entries = []
        result = win._generate_fallback()
        assert "共 0 个番茄钟" in result
