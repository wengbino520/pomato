from datetime import date, timedelta
import re

from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QFont, QTextCursor
from PyQt6.QtWidgets import (
    QApplication,
    QComboBox,
    QDialog,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QMessageBox,
    QProgressBar,
    QPushButton,
    QTextEdit,
    QVBoxLayout,
)

from src.services.logger import get_logger
from src.ui.styles import btn_style

logger = get_logger(__name__)


def _get_period_range(dt: date, period: str) -> tuple[date, date]:
    """Return (start, end) dates for the given period containing dt."""
    if period == "daily":
        return dt, dt
    elif period == "weekly":
        monday = dt - timedelta(days=dt.weekday())
        return monday, monday + timedelta(days=6)
    elif period == "monthly":
        first = dt.replace(day=1)
        if dt.month == 12:
            last = dt.replace(year=dt.year + 1, month=1, day=1) - timedelta(days=1)
        else:
            last = dt.replace(month=dt.month + 1, day=1) - timedelta(days=1)
        return first, last
    raise ValueError(f"Unknown period: {period}")


class _AIWorker(QThread):
    chunk_received = pyqtSignal(str)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, ai_client, entries: list[dict], report_date: str,
                 todos: list[dict] | None = None, period: str = "daily"):
        super().__init__()
        self.ai_client = ai_client
        self.entries = entries
        self.report_date = report_date
        self.todos = todos or []
        self.period = period

    def run(self):
        try:
            result = self.ai_client.generate_report(
                self.entries,
                self.report_date,
                on_chunk=lambda c: self.chunk_received.emit(c),
                todos=self.todos if self.todos else None,
                period=self.period,
            )
            self.finished.emit(result)
        except Exception as exc:
            self.error.emit(str(exc))


PERIOD_NAMES = {"daily": "日报", "weekly": "周报", "monthly": "月报"}


class ReportWindow(QDialog):
    def __init__(self, config, db, ai_client, parent=None, report_date=None,
                 period: str = "daily"):
        super().__init__(parent)
        self.config = config
        self.db = db
        self.ai_client = ai_client
        self.report_date = report_date or date.today().isoformat()
        self._period = period
        self._dt = date.fromisoformat(self.report_date)
        self._start_date, self._end_date = _get_period_range(self._dt, self._period)
        self.entries = self._load_entries()
        self._worker: _AIWorker | None = None
        self._setup_ui()
        # 从明确入口打开时隐藏周期下拉框（入口名称即周期）
        self.period_combo.hide()
        self._start_generation()

    # ------------------------------------------------------------------
    # UI
    # ------------------------------------------------------------------

    def _setup_ui(self):
        period_label = PERIOD_NAMES.get(self._period, "日报")
        self.setWindowTitle(f"POMATO · 工作{period_label}")
        self.setWindowFlags(Qt.WindowType.Window | Qt.WindowType.WindowStaysOnTopHint)
        self.resize(720, 620)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(10)

        # header
        hl = QHBoxLayout()
        title = QLabel(f"📋 生成报告")
        title.setStyleSheet("font-size:15px; font-weight:bold; color:#333;")

        # Period selector
        self.period_combo = QComboBox()
        self.period_combo.addItems(["日报", "周报", "月报"])
        self.period_combo.setCurrentIndex(0)
        self.period_combo.setStyleSheet(
            "QComboBox { background:#fff; border:1px solid #ccc; border-radius:4px;"
            "  padding:2px 6px; font-size:12px; }"
            "QComboBox::drop-down { border:none; }"
        )
        self.period_combo.currentIndexChanged.connect(self._on_period_changed)

        # Date range label
        self.date_range_label = QLabel(self._make_range_label())
        self.date_range_label.setStyleSheet("color:#757575; font-size:12px;")

        self.status_label = QLabel("AI 生成中…")
        self.status_label.setStyleSheet("color:#ef5350; font-size:12px;")

        hl.addWidget(title)
        hl.addWidget(self.period_combo)
        hl.addWidget(self.date_range_label)
        hl.addStretch()
        hl.addWidget(self.status_label)
        layout.addLayout(hl)

        # editor
        self.editor = QTextEdit()
        self.editor.setFont(QFont("Consolas", 12))
        self.editor.setStyleSheet(
            "QTextEdit { border:1px solid #eee; border-radius:6px;"
            "  padding:12px; background:#fafafa; }"
        )
        layout.addWidget(self.editor, 1)

        # indeterminate progress bar
        self.progress = QProgressBar()
        self.progress.setRange(0, 0)
        self.progress.setFixedHeight(4)
        self.progress.setStyleSheet(
            "QProgressBar { border:none; background:#eee; border-radius:2px; }"
            "QProgressBar::chunk { background:#ef5350; border-radius:2px; }"
        )
        layout.addWidget(self.progress)

        # buttons
        bl = QHBoxLayout()

        self.regenerate_btn = QPushButton("🔄 重新生成")
        self.regenerate_btn.setEnabled(False)
        self.regenerate_btn.setStyleSheet(btn_style("#757575", padding="8px 16px"))
        self.regenerate_btn.clicked.connect(self._start_generation)

        self.copy_btn = QPushButton("📋 复制")
        self.copy_btn.setEnabled(False)
        self.copy_btn.setStyleSheet(btn_style("#1976d2", padding="8px 16px"))
        self.copy_btn.clicked.connect(self._copy_to_clipboard)

        self.export_btn = QPushButton("💾 导出 Markdown")
        self.export_btn.setEnabled(False)
        self.export_btn.setStyleSheet(btn_style("#388e3c", padding="8px 16px"))
        self.export_btn.clicked.connect(self._export_markdown)

        self.export_docx_btn = QPushButton("📄 导出 Word")
        self.export_docx_btn.setEnabled(False)
        self.export_docx_btn.setStyleSheet(btn_style("#1565c0", padding="8px 16px"))
        self.export_docx_btn.clicked.connect(self._export_docx)

        close_btn = QPushButton("关闭")
        close_btn.setStyleSheet(btn_style("#9e9e9e", padding="8px 16px"))
        close_btn.clicked.connect(self.accept)

        bl.addWidget(self.regenerate_btn)
        bl.addStretch()
        bl.addWidget(self.copy_btn)
        bl.addWidget(self.export_btn)
        bl.addWidget(self.export_docx_btn)
        bl.addWidget(close_btn)
        layout.addLayout(bl)

    # ------------------------------------------------------------------
    # Period helpers
    # ------------------------------------------------------------------

    PERIOD_MAP = {"日报": "daily", "周报": "weekly", "月报": "monthly"}

    def _make_range_label(self) -> str:
        if self._start_date == self._end_date:
            return self._start_date.strftime("%m/%d")
        return f"{self._start_date.strftime('%m/%d')} ~ {self._end_date.strftime('%m/%d')}"

    def _load_entries(self):
        if self._start_date == self._end_date:
            return self.db.get_entries_by_date(self._start_date.isoformat())
        return self.db.get_entries_by_date_range(
            self._start_date.isoformat(), self._end_date.isoformat()
        )

    def _on_period_changed(self, index: int):
        period_key = ["日报", "周报", "月报"][index]
        self._period = self.PERIOD_MAP[period_key]
        self._start_date, self._end_date = _get_period_range(self._dt, self._period)
        self.date_range_label.setText(self._make_range_label())
        self.entries = self._load_entries()

    # ------------------------------------------------------------------
    # Generation flow
    # ------------------------------------------------------------------

    def _start_generation(self):
        self.editor.clear()
        self.progress.show()
        self.status_label.setText("AI 生成中…")
        self.status_label.setStyleSheet("color:#ef5350; font-size:12px;")
        self.regenerate_btn.setEnabled(False)
        self.copy_btn.setEnabled(False)
        self.export_btn.setEnabled(False)
        self.export_docx_btn.setEnabled(False)

        # Query todos by date range for weekly/monthly
        todos = self.db.get_todos_by_date_range(
            self._start_date.isoformat(), self._end_date.isoformat()
        )

        self._worker = _AIWorker(self.ai_client, self.entries, self.report_date,
                                 todos=todos, period=self._period)
        self._worker.chunk_received.connect(self._on_chunk)
        self._worker.finished.connect(self._on_finished)
        self._worker.error.connect(self._on_error)
        self._worker.start()

    def _on_chunk(self, chunk: str):
        cursor = self.editor.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.End)
        cursor.insertText(chunk)
        self.editor.setTextCursor(cursor)
        self.editor.ensureCursorVisible()

    def _on_finished(self, result: str):
        self.progress.hide()
        self.status_label.setText("✓ 生成完成，可直接编辑后导出")
        self.status_label.setStyleSheet("color:#388e3c; font-size:12px;")
        self.regenerate_btn.setEnabled(True)
        self.copy_btn.setEnabled(True)
        self.export_btn.setEnabled(True)
        self.export_docx_btn.setEnabled(True)
        # Always save (all periods persisted)
        try:
            self.db.save_report(self._start_date.isoformat(), self.entries,
                                period=self._period,
                                ai_summary=result, final_report=result)
            logger.info("%s report saved for %s (%d chars)",
                        self._period, self._start_date, len(result))
        except Exception:
            logger.exception("Failed to save %s report for %s",
                           self._period, self._start_date)

    def _on_error(self, error_msg: str):
        self.progress.hide()
        logger.error("AI report generation failed for %s: %s", self.report_date, error_msg)
        self.status_label.setText("⚠ AI 生成失败，已展示原始记录")
        self.status_label.setStyleSheet("color:#f44336; font-size:12px;")
        self.regenerate_btn.setEnabled(True)
        fallback = self._generate_fallback()
        self.editor.setPlainText(fallback)
        self.copy_btn.setEnabled(True)
        self.export_btn.setEnabled(True)
        self.export_docx_btn.setEnabled(True)
        QMessageBox.warning(
            self,
            "AI 生成失败",
            f"原因：{error_msg}\n\n已展示原始记录，您可手动整理后导出。",
        )

    def _generate_fallback(self) -> str:
        valid = [e for e in self.entries if not e.get("skipped") and e.get("content")]
        period_labels = {"daily": "今日", "weekly": "本周", "monthly": "本月"}
        label = period_labels.get(self._period, "今日")

        lines = [f"# 工作报告 · {self._make_range_label()}", ""]
        lines.append(f"## {label}记录（共 {len(valid)} 个番茄钟）")
        for e in valid:
            tags = ", ".join(e.get("tags") or [])
            tag_str = f"[{tags}] " if tags else ""
            date_prefix = ""
            if self._period != "daily":
                date_prefix = f"{e.get('date', '')} "
            lines.append(f"- {date_prefix}{e['start_time'][:5]}-{e['end_time'][:5]} {tag_str}{e['content']}")
        return "\n".join(lines)

    # ------------------------------------------------------------------
    # Actions
    # ------------------------------------------------------------------

    def _copy_to_clipboard(self):
        QApplication.clipboard().setText(self.editor.toPlainText())
        original = self.copy_btn.text()
        self.copy_btn.setText("✓ 已复制")
        from PyQt6.QtCore import QTimer
        QTimer.singleShot(2000, lambda: self.copy_btn.setText(original))

    def _export_markdown(self):
        period_names = {"daily": "日报", "weekly": "周报", "monthly": "月报"}
        default_name = f"{period_names.get(self._period, '报告')}_{self._make_range_label().replace(' ~ ', '-').replace('/', '-')}.md"
        path, _ = QFileDialog.getSaveFileName(
            self, "导出日报", default_name, "Markdown 文件 (*.md);;文本文件 (*.txt)"
        )
        if path:
            markdown = self.editor.toPlainText()
            text = markdown
            if path.lower().endswith(".txt"):
                text = self._markdown_to_plain_text(markdown)
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)
            self.db.save_report(self._start_date.isoformat(), self.entries,
                                period=self._period, final_report=markdown)
            QMessageBox.information(self, "导出成功", f"日报已保存至：\n{path}")

    @staticmethod
    def _markdown_to_plain_text(markdown: str) -> str:
        lines = []
        for line in markdown.splitlines():
            stripped = line.strip()
            if not stripped:
                lines.append("")
                continue

            stripped = re.sub(r"^#{1,6}\s*", "", stripped)
            stripped = re.sub(r"^[-*+]\s+", "- ", stripped)
            stripped = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            stripped = re.sub(r"\*(.*?)\*", r"\1", stripped)
            stripped = re.sub(r"`(.*?)`", r"\1", stripped)
            lines.append(stripped)

        plain = "\n".join(lines)
        plain = re.sub(r"\n{3,}", "\n\n", plain)
        return plain.strip() + "\n"

    def _export_docx(self):
        """将日报导出为 Word (.docx) 文件。"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            QMessageBox.warning(
                self, "缺少依赖",
                "导出 Word 需要 python-docx 库。\n请运行: pip install python-docx"
            )
            return

        period_names = {"daily": "日报", "weekly": "周报", "monthly": "月报"}
        default_name = f"{period_names.get(self._period, '报告')}_{self._make_range_label().replace(' ~ ', '-').replace('/', '-')}.docx"
        path, _ = QFileDialog.getSaveFileName(
            self, "导出 Word 日报", default_name,
            "Word 文档 (*.docx);;所有文件 (*)"
        )
        if not path:
            return

        markdown = self.editor.toPlainText()
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(11)

        for line in markdown.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            if stripped.startswith("# "):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped[2:])
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            elif stripped.startswith("## "):
                p = doc.add_paragraph()
                run = p.add_run(stripped[3:])
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0xE5, 0x39, 0x35)
            elif stripped.startswith("### "):
                p = doc.add_paragraph()
                run = p.add_run(stripped[4:])
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("> "):
                p = doc.add_paragraph()
                run = p.add_run(stripped[2:])
                run.italic = True
                run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
            else:
                # Handle inline bold (**text**) and italic (*text*)
                p = doc.add_paragraph()
                self._add_markdown_inline(p, stripped)

        doc.save(path)
        self.db.save_report(self._start_date.isoformat(), self.entries,
                            period=self._period, final_report=markdown)
        QMessageBox.information(self, "导出成功", f"日报已保存至：\n{path}")

    @staticmethod
    def _add_markdown_inline(paragraph, text: str):
        """将带 Markdown 内联格式的文本添加到 Word 段落。"""
        from docx.shared import Pt
        import re as _re

        # Split on **bold** and *italic* markers
        pattern = _re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")
        last_end = 0
        for m in pattern.finditer(text):
            # Text before the match
            if m.start() > last_end:
                paragraph.add_run(text[last_end:m.start()])
            if m.group(2):  # **bold**
                run = paragraph.add_run(m.group(2))
                run.bold = True
            elif m.group(3):  # *italic*
                run = paragraph.add_run(m.group(3))
                run.italic = True
            last_end = m.end()
        # Remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
